"""
Gerador de documentos Word (versão simples para testes)
NOTA: Versão definitiva será em .NET com DocumentFormat.OpenXml
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from models import ParsedDocument, ParsedNode, NodeType, TextType
from typing import Optional


class SimpleWordGenerator:
    """Gerador simples de Word para testes"""

    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        """Configura estilos básicos do documento"""
        # Estilo para título principal
        style = self.doc.styles['Heading 1']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(18)
        style.font.bold = True

        # Estilo para subtítulos
        style = self.doc.styles['Heading 2']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(14)
        style.font.bold = True

    def generate_from_parsed(self, parsed_doc: ParsedDocument, output_path: str):
        """
        Gera documento Word a partir de documento parseado

        Args:
            parsed_doc: Documento parseado e traduzido
            output_path: Caminho para salvar .docx
        """
        print(f"\n{'='*80}")
        print(f"GERANDO DOCUMENTO WORD")
        print(f"{'='*80}\n")

        # Título do documento
        title = self.doc.add_heading(parsed_doc.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar informação de origem
        info = self.doc.add_paragraph()
        info.add_run(f"Arquivo original: {parsed_doc.original_filename}\n").italic = True
        info.add_run("Traduzido automaticamente com IA").italic = True
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # Espaço

        # Processar nós
        nodes_processed = 0
        for node in parsed_doc.nodes:
            self._process_node(node)
            nodes_processed += 1

        # Salvar
        self.doc.save(output_path)

        print(f"[OK] Documento salvo: {output_path}")
        print(f"  - Nós processados: {nodes_processed}")
        print(f"{'='*80}\n")

    def _process_node(self, node: ParsedNode, level: int = 0):
        """Processa um nó recursivamente"""

        if node.node_type == NodeType.HEADING_2:
            self._add_heading(node, level=1)

        elif node.node_type == NodeType.HEADING_3:
            self._add_heading(node, level=2)

        elif node.node_type == NodeType.HEADING_4:
            self._add_heading(node, level=3)

        elif node.node_type == NodeType.PARAGRAPH:
            self._add_paragraph(node)

        elif node.node_type == NodeType.LIST_ORDERED:
            self._add_list(node, ordered=True)

        elif node.node_type == NodeType.LIST_UNORDERED:
            self._add_list(node, ordered=False)

        elif node.node_type == NodeType.TABLE:
            self._add_table(node)

        elif node.node_type == NodeType.BLOCKQUOTE:
            self._add_blockquote(node)

        # Processar filhos
        for child in node.children:
            self._process_node(child, level + 1)

    def _add_heading(self, node: ParsedNode, level: int):
        """Adiciona título"""
        text = self._get_node_text(node)
        if text:
            heading = self.doc.add_heading(text, level=level)

    def _add_paragraph(self, node: ParsedNode):
        """Adiciona parágrafo com formatação"""
        if not node.text_segments and not node.children:
            return

        para = self.doc.add_paragraph()

        # Adicionar número de seção se houver
        if node.section_number:
            run = para.add_run(f"{node.section_number}. ")
            run.bold = True

        # Adicionar segmentos de texto
        for segment in node.text_segments:
            run = para.add_run(segment.text)

            # Aplicar formatação
            if segment.formatting.bold:
                run.bold = True
            if segment.formatting.italic:
                run.italic = True

            # Destacar latim (opcional)
            if segment.text_type == TextType.LATIN:
                run.italic = True  # Latim sempre em itálico

        # Se tiver filhos inline, processar
        for child in node.children:
            if child.node_type in [NodeType.STRONG, NodeType.EMPHASIS]:
                self._add_inline_to_paragraph(para, child)

    def _add_inline_to_paragraph(self, para, node: ParsedNode):
        """Adiciona conteúdo inline a um parágrafo existente"""
        for segment in node.text_segments:
            run = para.add_run(segment.text)

            if node.node_type == NodeType.STRONG:
                run.bold = True
            elif node.node_type == NodeType.EMPHASIS:
                run.italic = True

            # Aplicar formatação adicional
            if segment.formatting.bold:
                run.bold = True
            if segment.formatting.italic:
                run.italic = True

    def _add_list(self, node: ParsedNode, ordered: bool = False):
        """Adiciona lista"""
        style = 'List Number' if ordered else 'List Bullet'

        for child in node.children:
            if child.node_type == NodeType.LIST_ITEM:
                text = self._get_node_text(child)
                if text:
                    para = self.doc.add_paragraph(text, style=style)

    def _add_table(self, node: ParsedNode):
        """Adiciona tabela simples"""
        # Contar linhas e colunas
        rows = [child for child in node.children if child.node_type == NodeType.TABLE_ROW]
        if not rows:
            return

        # Pegar número de colunas da primeira linha
        first_row = rows[0]
        cols = len([c for c in first_row.children if c.node_type in [NodeType.TABLE_CELL, NodeType.TABLE_HEADER]])

        if cols == 0:
            return

        # Criar tabela
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Light Grid Accent 1'

        # Preencher células
        for i, row_node in enumerate(rows):
            cells = [c for c in row_node.children if c.node_type in [NodeType.TABLE_CELL, NodeType.TABLE_HEADER]]
            for j, cell_node in enumerate(cells):
                if j < cols:
                    text = self._get_node_text(cell_node)
                    table.rows[i].cells[j].text = text

                    # Se for header, deixar em negrito
                    if cell_node.node_type == NodeType.TABLE_HEADER:
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _add_blockquote(self, node: ParsedNode):
        """Adiciona citação/nota"""
        text = self._get_node_text(node)
        if text:
            para = self.doc.add_paragraph(text)
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.right_indent = Inches(0.5)
            for run in para.runs:
                run.italic = True

    def _get_node_text(self, node: ParsedNode) -> str:
        """Extrai texto completo de um nó"""
        texts = []

        for segment in node.text_segments:
            texts.append(segment.text)

        for child in node.children:
            child_text = self._get_node_text(child)
            if child_text:
                texts.append(child_text)

        return ' '.join(texts).strip()
